from docx import Document
from docx.shared import Pt, RGBColor

import os
import random

class User:
    def __init__(self, username, password):
        self.username = username
        self.password = password

class Admin:
    def __init__(self, username, password):
        self.username = username
        self.password = password

def register(user_count):
    username = input("Kullanıcı adınızı girin: ")
    password = input("Şifrenizi girin: ")
    user = User(username, password)
    with open(f"kullanici{user_count}.txt", "w") as file:
        file.write(f"{user.username}\n{user.password}")
    print(f"Yeni kullanıcı başarıyla oluşturuldu ve 'kullanici{user_count}.txt' olarak kaydedildi: {user.username}")

def login(user_count):
    username = input("Kullanıcı adınızı girin: ")
    password = input("Şifrenizi girin: ")
    for i in range(1, user_count):
        with open(f"kullanici{i}.txt", "r") as file:
            lines = file.readlines()
            if username == lines[0].strip() and password == lines[1].strip():
                print(f"Başarıyla giriş yaptınız! Hoşgeldin {username}.")
                return username
    print("Hatalı kullanıcı adı veya şifre.")
    return None

def admin_login():
    username = input("Admin kullanıcı adınızı girin: ")
    password = input("Admin şifrenizi girin: ")
    if username == "admin" and password == "admin":
        print("Başarıyla admin girişi yaptınız!")
        return True
    print("Hatalı admin kullanıcı adı veya şifre.")
    return False

def add_word():
    english = input("Eklemek istediğiniz İngilizce kelimeyi girin: ")
    with open("word_pairs.txt", "r") as file:
        lines = file.readlines()
        for line in lines:
            existing_english, _, _ = line.strip().split(',')
            if english == existing_english:
                print(f"'{english}' kelimesi zaten mevcut.")
                return
    turkish = input("Eklemek istediğiniz kelimenin Türkçe karşılığını girin: ")
    with open("word_pairs.txt", "a") as file:
        file.write(f"{english},{turkish},0\n")
    print(f"'{english}' kelimesi başarıyla eklendi.")

def remove_word():
    word_to_remove = input("Silmek istediğiniz kelimeyi girin: ")
    lines_to_keep = []
    found = False
    with open("word_pairs.txt", "r") as file:
        lines = file.readlines()
        for line in lines:
            english, _, _ = line.strip().split(',')
            if english != word_to_remove:
                lines_to_keep.append(line)
            else:
                found = True
    if found:
        with open("word_pairs.txt", "w") as file:
            file.writelines(lines_to_keep)
        print(f"'{word_to_remove}' kelimesi başarıyla silindi.")
    else:
        print("Silmek istediğiniz kelime bulunamadı.")

def admin_screen():
    while True:
        print("1. Kelime Ekle")
        print("2. Kelime Sil")
        print("3. Ana Ekrana Dön")
        choice = input("Seçiminizi yapın (1-3): ")
        if choice == "1":
            add_word()
        elif choice == "2":
            remove_word()
        elif choice == "3":
            print("Ana ekrana dönülüyor.")
            break


def main_screen():
    print("1. Giriş Yap")
    print("2. Kayıt Ol")
    print("3. Admin Girişi")
    choice = input("Seçiminizi yapın (1-3): ")
    return choice

def english_learning_module(user):
    print("1. Pratik Yap")
    print("2. Sınava Gir")
    choice = input("Ne yapmak istersiniz? (1-2): ")
    word_pairs = {}
    with open("word_pairs.txt", "r") as file:
        lines = file.readlines()
        for line in lines:
            english, turkish, count = line.strip().split(',')
            if int(count) < 6:
                word_pairs[english] = [turkish, int(count)]

    if not word_pairs:
        print("Uygun kelime çifti kalmadı. Hepsini öğrendin.")
        return

    if choice == "1":
        num_words = int(input("Kaç kelime ile pratik yapmak istersiniz? (0-50): "))
        num_words = min(max(num_words, 0), 50)
        correct_count = 0
        for _ in range(num_words):
            min_count = min(word_pairs.values(), key=lambda x: x[1])[1]
            words_with_min_count = [word for word in word_pairs if word_pairs[word][1] == min_count]
            english_word = random.choice(words_with_min_count)
            turkish_word = word_pairs[english_word][0]
            print(f"Bu İngilizce kelimenin Türkçesi nedir? {english_word}")
            for i in range(2):
                user_answer = input("Cevabınız: ")
                if user_answer == turkish_word:
                    print("Doğru cevap!")
                    word_pairs[english_word][1] += 1
                    correct_count += 1
                    if word_pairs[english_word][1] == 6:
                        del word_pairs[english_word]
                    break
                elif i == 0:
                    print("Yanlış cevap, tekrar dene.")
                else:
                    print(f"Üzgünüm, bilemediniz. Cevap {turkish_word} olacaktı.")
        print(f"Toplam {num_words} soru soruldu. {correct_count} tanesini doğru, {num_words - correct_count} tanesini yanlış yanıtladınız.")
    elif choice == "2":
        num_words = 20
        correct_count = 0
        exam_words = random.sample(list(word_pairs.keys()), num_words)
        exam_report = []
        for i in range(num_words):
            english_word = exam_words[i]
            turkish_word = word_pairs[english_word][0]
            print(f"{i+1}. soru: Bu İngilizce kelimenin Türkçesi nedir? {english_word}")
            user_answer = input("Cevabınız: ")
            score = 5 if user_answer == turkish_word else 0
            correct_count += score
            exam_report.append(f"{i+1}. soru: {english_word}, Cevabınız: {user_answer}, Puan: {score}")
        total_score = correct_count
        print(f"Sınavdan {total_score}/100 puan aldınız.")
        report_choice = input("Rapor ister misiniz? (y/n): ")
        if report_choice.lower() == "y":
            doc = Document()
            title = doc.add_heading(level=1)
            run = title.add_run(f"{user} adlı kişinin sınav raporu: {total_score}")
            run.font.color.rgb = RGBColor(255, 0, 0)  # Kırmızı renk
            for line in exam_report:
                question, user_answer, score = line.split(", ")
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(question)
                run.bold = True  # Kalın yazı tipi
                run = paragraph.add_run(", " + user_answer)
                run.font.color.rgb = RGBColor(0, 255, 0) if int(score.split(": ")[1]) > 0 else RGBColor(255, 0, 0)
                paragraph.add_run(", " + score)
            doc.save("sinavraporu.docx")
            print("Raporunuz 'sinavraporu.docx' olarak kaydedildi.")

    with open("word_pairs.txt", "w") as file:
        for english, (turkish, count) in word_pairs.items():
            file.write(f"{english},{turkish},{count}\n")

user_count = len([name for name in os.listdir() if name.startswith("kullanici") and name.endswith(".txt")]) + 1
while True:
    choice = main_screen()
    if choice == "1":
        user = login(user_count)
        if user:
            english_learning_module(user)
    elif choice == "2":
        register(user_count)
        user_count += 1
    elif choice == "3":
        if admin_login():
            admin_screen()
